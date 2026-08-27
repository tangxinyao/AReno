"""Reproducible Office fixtures, oracle artifacts, and deterministic grading."""

from __future__ import annotations

import json
import random
import re
import string
import tempfile
from pathlib import Path
from typing import Any

DEFAULT_TEMPLATE_DIR = Path(__file__).with_name("templates")


def load_template_catalog(template_dir: Path | str = DEFAULT_TEMPLATE_DIR) -> dict[str, dict[str, Any]]:
    """Load and validate the structured Office template catalog."""

    path = Path(template_dir) / "catalog.json"
    payload = json.loads(path.read_text(encoding="utf-8"))
    if payload.get("schema_version") != 1:
        raise ValueError(f"unsupported Office template schema in {path}")
    templates = payload.get("templates")
    if not isinstance(templates, list) or not templates:
        raise ValueError(f"Office template catalog has no templates: {path}")
    catalog = {}
    for template in templates:
        task = template.get("task")
        if not isinstance(task, str) or task in catalog:
            raise ValueError(f"invalid or duplicate Office task in {path}: {task!r}")
        if not template.get("prompt_templates") or not isinstance(template.get("fields"), dict):
            raise ValueError(f"incomplete Office template in {path}: {task}")
        catalog[task] = template
    return catalog


_DEFAULT_CATALOG = load_template_catalog()
TASK_TYPES = tuple(_DEFAULT_CATALOG)


def _random_token(rng: random.Random, *, minimum: int = 9, maximum: int = 15) -> str:
    length = rng.randint(minimum, maximum)
    first = rng.choice(string.ascii_lowercase)
    rest = "".join(rng.choices(string.ascii_lowercase + string.digits, k=length - 1))
    return first + rest


_HAN_ALPHABET = "安博辰达恩丰格禾景凯朗明宁启然森拓维新彦云泽春川峰海湖林岭秋山石松田星雪阳原月舟竹"


def _random_text(rng: random.Random, *, minimum: int = 2, maximum: int = 6) -> str:
    """Generate deterministic free text from characters, not a value enum."""

    return "".join(rng.choices(_HAN_ALPHABET, k=rng.randint(minimum, maximum)))


def _unique_random_texts(
    rng: random.Random,
    count: int,
    *,
    minimum: int = 2,
    maximum: int = 6,
) -> list[str]:
    values: list[str] = []
    while len(values) < count:
        value = _random_text(rng, minimum=minimum, maximum=maximum)
        if value not in values:
            values.append(value)
    return values


def _generate_field(rng: random.Random, definition: dict[str, Any]) -> str:
    generator = definition.get("generator")
    if generator == "filename":
        extension = definition.get("extension")
        if not isinstance(extension, str) or not extension.startswith("."):
            raise ValueError(f"filename generator requires an extension: {definition!r}")
        return f"{_random_token(rng)}{extension}"
    if generator == "text":
        return _random_text(
            rng,
            minimum=int(definition.get("min_length", 2)),
            maximum=int(definition.get("max_length", 6)),
        )
    raise ValueError(f"unsupported Office field generator: {generator!r}")


def build_record(
    task: str,
    *,
    seed: int,
    index: int,
    template_dir: Path | str = DEFAULT_TEMPLATE_DIR,
) -> dict[str, Any]:
    """Build one self-contained task description from a deterministic seed."""

    catalog = _DEFAULT_CATALOG if Path(template_dir) == DEFAULT_TEMPLATE_DIR else load_template_catalog(template_dir)
    if task not in catalog:
        raise ValueError(f"unsupported Office task: {task}")
    rng = random.Random(seed)
    template = catalog[task]
    generated_fields = {name: _generate_field(rng, definition) for name, definition in template["fields"].items()}
    output_file = generated_fields["output_file"]
    products = _unique_random_texts(rng, 12, minimum=2, maximum=5)
    rows = [
        {
            "name": name,
            "price": rng.randint(20, 300),
            "quantity": rng.randint(2, 20),
        }
        for name in products[:8]
    ]
    monthly = [
        {"month": f"{month}月", "sales": rng.randint(80, 260), "revenue": rng.randint(1800, 7200)}
        for month in range(1, rng.randint(4, 10))
    ]
    discount_rows = [
        {"name": row["name"], "original_price": row["price"], "discount_rate": rng.choice([0.05, 0.1, 0.15, 0.2])}
        for row in rows[: rng.randint(4, 9)]
    ]
    region_pool = _unique_random_texts(rng, 5, minimum=2, maximum=4)
    regions = region_pool[: rng.randint(2, len(region_pool) + 1)]
    regional_rows = [
        {"name": row["name"], "regions": {region: rng.randint(800, 5000) for region in regions}}
        for row in rows[: rng.randint(3, len(rows) + 1)]
    ]
    vocabulary_pool = _unique_random_texts(rng, 8, minimum=2, maximum=4)
    vocabulary = vocabulary_pool[: rng.randint(3, 7)]
    vocabulary_counts = {term: rng.randint(1, 4) for term in vocabulary}
    distractors = _unique_random_texts(rng, rng.randint(1, 4), minimum=2, maximum=4)
    document_title = _random_text(rng, minimum=4, maximum=8)
    document_paragraph_pool = [
        f"{_random_text(rng, minimum=2, maximum=5)}已完成{_random_text(rng, minimum=2, maximum=5)}，"
        f"{_random_text(rng, minimum=2, maximum=5)}正在推进。"
        for _ in range(5)
    ]
    spec = {
        "task": task,
        "template_name": Path(template_dir).name,
        "risk_tags": template.get("risk_tags", []),
        **generated_fields,
        "products": rows,
        "monthly": monthly,
        "discount_rows": discount_rows,
        "regions": regions,
        "regional_rows": regional_rows,
        "vocabulary": vocabulary,
        "vocabulary_counts": vocabulary_counts,
        "distractors": distractors,
        "document_title": document_title,
        "document_paragraphs": document_paragraph_pool[: rng.randint(2, 6)],
    }
    return {
        "id": f"office-{index:05d}",
        "task": task,
        "seed": seed,
        "task_spec": json.dumps(spec, ensure_ascii=False, separators=(",", ":")),
        "output_file": output_file,
        "max_turns": 12,
        "context_budget": 10_000,
        "prompt": make_prompt(spec, template=template, rng=rng),
    }


def make_prompt(
    spec: dict[str, Any],
    *,
    template: dict[str, Any] | None = None,
    rng: random.Random | None = None,
) -> str:
    """Render a task through a compositional grammar, not one fixed pattern."""

    selected = template or _DEFAULT_CATALOG[spec["task"]]
    prompt_rng = rng or random.Random(int(spec.get("seed", 0)))
    body = prompt_rng.choice(selected["prompt_templates"]).format_map(spec)
    clauses = [part.strip() for part in re.split(r"[。；]", body) if part.strip()]
    prefix = prompt_rng.choice(
        (
            "请完成下面的 Office 任务：",
            "在当前工作目录中处理以下交付：",
            "任务说明：",
            "请按要求生成并核验交付物：",
            "需要完成的文档操作如下：",
            "",
        )
    )
    suffix = prompt_rng.choice(
        (
            "完成后请检查目标文件确实存在且能够重新打开。",
            "请使用相对路径，并以最终保存的文件作为交付结果。",
            "不要只描述步骤；请实际生成文件并进行一次结果核验。",
            "验收以目标文件中的实际内容和对象为准。",
            "",
        )
    )
    layout = prompt_rng.randrange(6)
    if layout == 0:
        pieces = [prefix, body, suffix]
        return " ".join(piece for piece in pieces if piece)
    if layout == 1:
        steps = "\n".join(f"{index}. {clause}。" for index, clause in enumerate(clauses, start=1))
        return "\n".join(piece for piece in (prefix, steps, suffix) if piece)
    if layout == 2:
        bullets = "\n".join(f"- {clause}" for clause in clauses)
        return "\n".join(piece for piece in (prefix, "要求：", bullets, suffix) if piece)
    if layout == 3:
        head, *rest = clauses
        details = "；".join(rest)
        pieces = [prefix, f"交付目标：{head}。"]
        if details:
            pieces.append(f"处理约束：{details}。")
        if suffix:
            pieces.append(f"验收：{suffix}")
        return "\n".join(piece for piece in pieces if piece)
    if layout == 4:
        steps = " / ".join(clauses)
        pieces = [prefix, f"操作链：{steps}。", suffix]
        return "\n\n".join(piece for piece in pieces if piece)
    lines = [prefix, "[任务]"] if prefix else ["[任务]"]
    lines.extend(f"[{index}] {clause}" for index, clause in enumerate(clauses, start=1))
    if suffix:
        lines.extend(("[验收]", suffix))
    return "\n".join(lines)


def parse_spec(record: dict[str, Any]) -> dict[str, Any]:
    value = record.get("task_spec")
    if isinstance(value, str):
        parsed = json.loads(value)
    elif isinstance(value, dict):
        parsed = dict(value)
    else:
        raise ValueError("Office record must contain task_spec")
    return parsed


def materialize_fixtures(record: dict[str, Any], root: Path) -> None:
    """Create binary input fixtures in an isolated rollout workspace."""

    spec = parse_spec(record)
    task = spec["task"]
    root.mkdir(parents=True, exist_ok=True)
    if task == "excel_chart_dashboard":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = spec["sheet_name"]
        sheet.append(["月份", "销售量", "营收"])
        for row in spec["monthly"]:
            sheet.append([row["month"], row["sales"], row["revenue"]])
        workbook.save(root / spec["input_file"])
    elif task == "excel_formula_conditional":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = spec["sheet_name"]
        sheet.append(["产品名称", "原价", "折扣率"])
        for row in spec["discount_rows"]:
            sheet.append([row["name"], row["original_price"], row["discount_rate"]])
        workbook.save(root / spec["input_file"])
    elif task == "office_mixed_report":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = spec["sheet_name"]
        sheet.append(["产品", *spec["regions"], "合计"])
        for row in spec["regional_rows"]:
            values = [row["regions"][region] for region in spec["regions"]]
            sheet.append([row["name"], *values, sum(values)])
        totals = [sum(row["regions"][region] for row in spec["regional_rows"]) for region in spec["regions"]]
        sheet.append(["合计", *totals, sum(totals)])
        workbook.save(root / spec["input_file"])
    elif task == "excel_read_analyze":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = spec["sheet_name"]
        sheet.append(["产品", *spec["regions"], "合计"])
        for row in spec["regional_rows"]:
            values = [row["regions"][region] for region in spec["regions"]]
            sheet.append([row["name"], *values, sum(values)])
        totals = [sum(row["regions"][region] for row in spec["regional_rows"]) for region in spec["regions"]]
        sheet.append(["合计", *totals, sum(totals)])
        workbook.save(root / spec["input_file"])
    elif task == "word_table_doc":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = spec["sheet_name"]
        sheet.append(["月份", "销售量"])
        for row in spec["monthly"]:
            sheet.append([row["month"], row["sales"]])
        workbook.save(root / spec["input_file"])
    elif task == "word_convert_pdf":
        from docx import Document

        document = Document()
        document.add_heading(spec["document_title"], level=0)
        for paragraph in spec["document_paragraphs"]:
            document.add_paragraph(paragraph)
        document.save(root / spec["input_file"])
    elif task == "word_extract_stat":
        from docx import Document

        document = Document()
        tokens = [term for term in spec["vocabulary"] for _ in range(spec["vocabulary_counts"][term])] + list(
            spec["distractors"]
        )
        random.Random(int(record["seed"])).shuffle(tokens)
        split = max(1, len(tokens) // 2)
        document.add_paragraph("，".join(tokens[:split]) + "。")
        document.add_paragraph("，".join(tokens[split:]) + "。")
        document.save(root / spec["input_file"])
        (root / spec["vocabulary_file"]).write_text("\n".join(spec["vocabulary"]) + "\n", encoding="utf-8")
    elif task == "word_format_edit":
        from docx import Document

        document = Document()
        document.add_heading(spec["document_title"], level=1)
        for paragraph in spec["document_paragraphs"]:
            document.add_paragraph(paragraph)
        document.add_paragraph("三、待协调事项：需要确认资源安排。")
        document.save(root / spec["input_file"])


def grade_workspace(record: dict[str, Any], root: Path) -> dict[str, Any]:
    """Grade the generated artifact without an LLM judge."""

    spec = parse_spec(record)
    output = root / spec["output_file"]
    if not output.is_file() or output.stat().st_size == 0:
        return {"score": 0.0, "issues": [f"missing output: {output.name}"]}
    try:
        if output.suffix == ".xlsx":
            issues = _grade_xlsx(spec, output)
        elif output.suffix == ".docx":
            issues = _grade_docx(spec, output)
        elif output.suffix == ".pdf":
            issues = _grade_pdf(output)
        elif output.suffix == ".json":
            issues = _grade_word_frequency(spec, output)
        else:
            issues = _grade_summary(spec, output)
    except Exception as exc:
        return {"score": 0.0, "issues": [f"cannot parse output: {exc}"]}
    return {"score": max(0.0, 1.0 - 0.25 * len(issues)), "issues": issues}


def _grade_xlsx(spec: dict[str, Any], path: Path) -> list[str]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=False)
    sheet = workbook.active
    issues = []
    task = spec["task"]
    if task == "excel_chart_dashboard":
        expected = [[row["month"], row["sales"], row["revenue"]] for row in spec["monthly"]]
        actual = [[sheet.cell(r, c).value for c in range(1, 4)] for r in range(2, 2 + len(expected))]
        if actual != expected:
            issues.append("dashboard data does not match fixture")
        if not sheet._charts:
            issues.append("dashboard has no chart")
    elif task == "excel_formula_conditional":
        if sheet.max_row < len(spec["discount_rows"]) + 2:
            issues.append("missing product or total rows")
        for row_index in range(2, 2 + len(spec["discount_rows"])):
            if not str(sheet.cell(row_index, 4).value or "").startswith("="):
                issues.append("discount price column must contain formulas")
                break
        if not sheet.conditional_formatting:
            issues.append("missing conditional formatting")
        if "合计" not in {sheet.cell(sheet.max_row, column).value for column in range(1, 6)}:
            issues.append("missing total row")
    else:
        headers = [sheet.cell(1, column).value for column in range(1, 5)]
        if headers != ["产品名称", "单价", "数量", "金额"]:
            issues.append("incorrect quote headers")
        if not all(sheet.cell(1, column).font.bold for column in range(1, 5)):
            issues.append("quote headers are not bold")
        if any((sheet.column_dimensions[letter].width or 0) <= 0 for letter in "ABCD"):
            issues.append("quote column widths are not set")
        if sheet.max_row != 9:
            issues.append("quote must contain exactly eight data rows")
        for row_index in range(2, 10):
            name, price, quantity = [sheet.cell(row_index, column).value for column in range(1, 4)]
            if (
                not name
                or not isinstance(price, int | float)
                or price <= 0
                or not isinstance(quantity, int | float)
                or quantity <= 0
            ):
                issues.append("quote rows require a name, positive price, and positive quantity")
                break
            if not str(sheet.cell(row_index, 4).value or "").startswith("="):
                issues.append("quote amount must use a formula")
                break
    workbook.close()
    return issues


def _grade_docx(spec: dict[str, Any], path: Path) -> list[str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    issues = []
    if spec["task"] == "word_create_report":
        for term in ("《工作周报》", "一、本周进展", "二、遇到的问题", "三、下周计划", "填写人："):
            if term not in text:
                issues.append(f"missing report term: {term}")
    elif spec["task"] == "word_format_edit":
        if spec["document_title"] not in text or "三、待协调事项" not in text or "填写人：" not in text:
            issues.append("formatted document lost source text")
        if not document.paragraphs or document.paragraphs[0].alignment != WD_ALIGN_PARAGRAPH.CENTER:
            issues.append("title is not centered")
        body_runs = [run for paragraph in document.paragraphs[1:] for run in paragraph.runs if run.text.strip()]
        if not body_runs or any(run.font.size != Pt(12) for run in body_runs):
            issues.append("body is not 12pt")
        if any((run.font.name or "") not in {"宋体", "SimSun"} for run in body_runs):
            issues.append("body is not SimSun")
        coordination = [paragraph for paragraph in document.paragraphs if "三、待协调事项" in paragraph.text]
        if not coordination or "List" not in coordination[0].style.name:
            issues.append("coordination paragraph is not a bullet")
    elif spec["task"] == "office_mixed_report":
        if not all(term in text for term in ("销售报告", "数据表", "结论")):
            issues.append("mixed report is missing required headings")
        if not document.tables:
            issues.append("mixed report has no table")
        expected_total = sum(sum(row["regions"].values()) for row in spec["regional_rows"])
        table_text = " ".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        if str(expected_total) not in f"{text} {table_text}":
            issues.append("mixed report has no correct revenue total")
        product_totals = {row["name"]: sum(row["regions"].values()) for row in spec["regional_rows"]}
        region_totals = {
            region: sum(row["regions"][region] for row in spec["regional_rows"]) for region in spec["regions"]
        }
        top_product = max(product_totals, key=product_totals.get)
        top_region = max(region_totals, key=region_totals.get)
        for value in (top_product, str(product_totals[top_product]), top_region, str(region_totals[top_region])):
            if value not in text:
                issues.append("mixed report conclusions do not match source data")
                break
    else:
        expected_rows = [["月份", "销售量"], *[[row["month"], str(row["sales"])] for row in spec["monthly"]]]
        if not document.tables:
            issues.append("monthly sales document has no table")
        else:
            table = document.tables[0]
            actual_rows = [[cell.text for cell in row.cells] for row in table.rows]
            if actual_rows != expected_rows:
                issues.append(f"monthly sales table does not match {spec['input_file']}")
            if not all(
                run.bold for cell in table.rows[0].cells for paragraph in cell.paragraphs for run in paragraph.runs
            ):
                issues.append("monthly sales headers are not bold")
        total = sum(row["sales"] for row in spec["monthly"])
        table_index = next(
            (index for index, paragraph in enumerate(document.paragraphs) if str(total) in paragraph.text), -1
        )
        if table_index < 0:
            issues.append("monthly sales total is missing below the table")
    return issues


def _grade_pdf(path: Path) -> list[str]:
    from pypdf import PdfReader

    content = path.read_bytes()
    issues = []
    if not content.startswith(b"%PDF-") or not content.rstrip().endswith(b"%%EOF"):
        issues.append("output is not a structurally complete PDF")
    try:
        reader = PdfReader(path)
        page_count = len(reader.pages)
    except Exception:
        page_count = 0
    if page_count < 1 or len(content) < 300:
        issues.append("PDF has no parseable page content")
    return issues


def _grade_word_frequency(spec: dict[str, Any], path: Path) -> list[str]:
    try:
        actual = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return ["word frequency output is not valid UTF-8 JSON"]
    expected = spec["vocabulary_counts"]
    return (
        []
        if actual == expected
        else [f"word frequencies do not match {spec['input_file']} and {spec['vocabulary_file']}"]
    )


def _grade_summary(spec: dict[str, Any], path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8", errors="replace")
    issues = []
    rows = spec["regional_rows"]
    product_totals = {row["name"]: sum(row["regions"].values()) for row in rows}
    overall = sum(product_totals.values())
    if str(len(rows)) not in text or "产品" not in text:
        issues.append("summary has no product count")
    if str(len(spec["regions"])) not in text or "区域" not in text:
        issues.append("summary has no region count")
    if any(name not in text or str(total) not in text for name, total in product_totals.items()):
        issues.append("summary has incorrect per-product revenue")
    if str(overall) not in text or "总营收" not in text:
        issues.append("summary has no correct overall revenue")
    if spec["input_file"] not in text or "来源" not in text:
        issues.append("summary does not identify its source")
    return issues


def create_oracle_artifact(record: dict[str, Any], root: Path) -> None:
    """Create a known-good artifact used to gate generated dataset records."""

    spec = parse_spec(record)
    output = root / spec["output_file"]
    if output.suffix == ".xlsx":
        _create_oracle_xlsx(spec, output)
    elif output.suffix == ".docx":
        _create_oracle_docx(spec, root, output)
    elif output.suffix == ".pdf":
        _create_oracle_pdf(output)
    elif output.suffix == ".json":
        output.write_text(json.dumps(spec["vocabulary_counts"], ensure_ascii=False), encoding="utf-8")
    else:
        _create_oracle_summary(spec, output)


def _create_oracle_xlsx(spec: dict[str, Any], output: Path) -> None:
    from openpyxl import Workbook, load_workbook
    from openpyxl.chart import BarChart, Reference
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import Font, PatternFill

    task = spec["task"]
    if task in {"excel_chart_dashboard", "excel_formula_conditional"}:
        workbook = load_workbook(output.parent / spec["input_file"])
        sheet = workbook.active
    else:
        workbook = Workbook()
        sheet = workbook.active
    if task == "excel_chart_dashboard":
        sheet.title = spec["sheet_name"]
        chart = BarChart()
        chart.add_data(Reference(sheet, min_col=3, min_row=1, max_row=sheet.max_row), titles_from_data=True)
        chart.set_categories(Reference(sheet, min_col=1, min_row=2, max_row=sheet.max_row))
        sheet.add_chart(chart, "E2")
    elif task == "excel_formula_conditional":
        sheet.cell(1, 4, "折扣价")
        for row_index in range(2, sheet.max_row + 1):
            sheet.cell(row_index, 4, f"=B{row_index}*(1-C{row_index})")
        total_row = sheet.max_row + 1
        sheet.cell(total_row, 1, "合计")
        sheet.cell(total_row, 2, f"=SUM(B2:B{total_row - 1})")
        sheet.cell(total_row, 4, f"=SUM(D2:D{total_row - 1})")
        average = sum(row["discount_rate"] for row in spec["discount_rows"]) / len(spec["discount_rows"])
        sheet.conditional_formatting.add(
            f"C2:C{total_row - 1}",
            CellIsRule(operator="greaterThan", formula=[str(average)], fill=PatternFill("solid", fgColor="00FF00")),
        )
    else:
        sheet.title = spec["sheet_name"]
        sheet.append(["产品名称", "单价", "数量", "金额"])
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for letter, width in {"A": 18, "B": 12, "C": 12, "D": 14}.items():
            sheet.column_dimensions[letter].width = width
        for row_index, row in enumerate(spec["products"], start=2):
            sheet.append([row["name"], row["price"], row["quantity"], f"=B{row_index}*C{row_index}"])
    workbook.save(output)


def _create_oracle_docx(spec: dict[str, Any], root: Path, output: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    task = spec["task"]
    if task == "word_format_edit":
        document = Document(root / spec["input_file"])
        document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in document.paragraphs[0].runs:
            run.bold = True
        document.add_paragraph("填写人：____________")
        for paragraph in document.paragraphs[1:]:
            for run in paragraph.runs:
                run.font.name = "SimSun"
                run.font.size = Pt(12)
            if "三、待协调事项" in paragraph.text:
                paragraph.style = "List Bullet"
    elif task == "word_create_report":
        document = Document()
        document.add_heading("《工作周报》", level=0)
        for heading, body in (
            ("一、本周进展", "完成数据整理和项目复盘。关键事项已按计划交付。"),
            ("二、遇到的问题", "部分数据仍需确认。相关风险已经同步负责人。"),
            ("三、下周计划", "完成验证并提交报告。继续跟进遗留事项。"),
        ):
            document.add_heading(heading, level=1)
            document.add_paragraph(body)
        document.add_paragraph("填写人：____")
    elif task == "office_mixed_report":
        document = Document()
        document.add_heading("销售报告", level=0)
        document.add_heading("数据表格", level=1)
        table = document.add_table(rows=1, cols=len(spec["regions"]) + 2)
        for cell, value in zip(table.rows[0].cells, ("产品", *spec["regions"], "合计"), strict=True):
            cell.text = value
        for row in spec["regional_rows"]:
            cells = table.add_row().cells
            values = [row["regions"][region] for region in spec["regions"]]
            for cell, value in zip(cells, (row["name"], *values, sum(values)), strict=True):
                cell.text = str(value)
        total = sum(sum(row["regions"].values()) for row in spec["regional_rows"])
        document.add_heading("结论", level=1)
        product_totals = {row["name"]: sum(row["regions"].values()) for row in spec["regional_rows"]}
        region_totals = {
            region: sum(row["regions"][region] for row in spec["regional_rows"]) for region in spec["regions"]
        }
        top_product = max(product_totals, key=product_totals.get)
        top_region = max(region_totals, key=region_totals.get)
        document.add_paragraph(
            f"整体营收合计：{total}。最高产品为{top_product}，营收{product_totals[top_product]}；"
            f"最高区域为{top_region}，营收{region_totals[top_region]}。"
        )
    else:
        document = Document()
        table = document.add_table(rows=1, cols=2)
        for cell, value in zip(table.rows[0].cells, ("月份", "销售量"), strict=True):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in spec["monthly"]:
            cells = table.add_row().cells
            cells[0].text = row["month"]
            cells[1].text = str(row["sales"])
        document.add_paragraph(f"所有销售量合计为：{sum(row['sales'] for row in spec['monthly'])}")
    document.save(output)


def _create_oracle_summary(spec: dict[str, Any], output: Path) -> None:
    rows = spec["regional_rows"]
    product_totals = {row["name"]: sum(row["regions"].values()) for row in rows}
    lines = [
        f"产品总数：{len(rows)}（来源：{spec['input_file']} 产品行）",
        f"区域总数：{len(spec['regions'])}（来源：{spec['input_file']} 区域列）",
        *[
            f"{name} 营收总额：{total}（来源：{spec['input_file']} 各区域求和）"
            for name, total in product_totals.items()
        ],
        f"整体总营收：{sum(product_totals.values())}（来源：{spec['input_file']} 产品合计）",
    ]
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _create_oracle_pdf(output: Path) -> None:
    """Write a small standards-compliant one-page PDF without optional dependencies."""

    stream = b"BT /F1 14 Tf 72 720 Td (Converted from sample.docx) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode())
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode())
    content.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode())
    output.write_bytes(content)


def validate_record(record: dict[str, Any]) -> dict[str, Any]:
    """Return the oracle grade for one generated record."""

    with tempfile.TemporaryDirectory(prefix="areno-office-oracle-") as directory:
        root = Path(directory)
        materialize_fixtures(record, root)
        create_oracle_artifact(record, root)
        return grade_workspace(record, root)
