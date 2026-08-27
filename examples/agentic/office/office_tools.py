"""Hermes-compatible tools backed by an isolated Office workspace."""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from office_env import grade_workspace, materialize_fixtures

MAX_OUTPUT_CHARS = 4_000

TOOLS = json.loads((Path(__file__).with_name("hermes_tool_schemas.json")).read_text(encoding="utf-8"))


_SKILLS = {
    "xlsx": (
        "Use openpyxl. Load inputs with load_workbook('sales.xlsx'), create formulas as strings beginning with '=', "
        "add charts with openpyxl.chart, save to the exact relative output filename, reopen it, and assert required "
        "cells/formulas/charts before finishing."
    ),
    "docx": (
        "Use python-docx. Load an input with Document('template.docx') or create Document(), use headings, paragraphs, "
        "tables, WD_ALIGN_PARAGRAPH and Pt for formatting, save to the exact relative output filename, reopen it, "
        "and assert required text/tables before finishing. Convert DOCX to PDF with "
        "subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '.', 'sample.docx'])."
    ),
}


class OfficeToolError(ValueError):
    """Compact model-visible workspace error."""


@dataclass(slots=True)
class OfficeWorkspace:
    record: dict[str, Any]
    root: Path
    milestones: set[str] = field(default_factory=set)

    @classmethod
    def from_record(cls, record: dict[str, Any]) -> OfficeWorkspace:
        root = Path(tempfile.mkdtemp(prefix="areno-office-"))
        try:
            materialize_fixtures(record, root)
            runtime = root / ".areno_runtime"
            runtime.mkdir()
            (runtime / "hermes_tools.py").write_text(_HERMES_TOOLS_COMPAT, encoding="utf-8")
            return cls(record=record, root=root)
        except Exception:
            shutil.rmtree(root, ignore_errors=True)
            raise

    def close(self) -> None:
        shutil.rmtree(self.root, ignore_errors=True)

    def execute_code(self, code: str) -> dict[str, Any]:
        if not code.strip():
            raise OfficeToolError("code must be non-empty")
        try:
            environment = dict(os.environ)
            runtime_path = str(self.root / ".areno_runtime")
            environment["PYTHONPATH"] = os.pathsep.join(
                value for value in (runtime_path, environment.get("PYTHONPATH")) if value
            )
            completed = subprocess.run(
                [sys.executable, "-c", code],
                cwd=self.root,
                text=True,
                capture_output=True,
                timeout=90,
                check=False,
                env=environment,
            )
            result = {
                "status": "success" if completed.returncode == 0 else "error",
                "returncode": completed.returncode,
                "stdout": completed.stdout[-MAX_OUTPUT_CHARS:],
                "stderr": completed.stderr[-MAX_OUTPUT_CHARS:],
            }
            if completed.returncode == 0:
                self.milestones.add("executed")
        except subprocess.TimeoutExpired as exc:
            result = {
                "status": "timeout",
                "returncode": 124,
                "stdout": str(exc.stdout or "")[-MAX_OUTPUT_CHARS:],
                "stderr": str(exc.stderr or "")[-MAX_OUTPUT_CHARS:],
            }
        return self._with_grade(result)

    def read_file(self, path: str, offset: int = 1, limit: int = 2_000) -> dict[str, Any]:
        target = _safe_path(self.root, path)
        if not target.is_file():
            raise OfficeToolError(f"not a file: {path}")
        if target.suffix.lower() == ".xlsx":
            content = _read_xlsx(target)
        elif target.suffix.lower() == ".docx":
            content = _read_docx(target)
        elif target.suffix.lower() == ".pdf":
            content = _read_pdf(target)
        else:
            content = target.read_text(encoding="utf-8", errors="replace")
        lines = content.splitlines()
        start = max(int(offset), 1) - 1
        count = min(max(int(limit), 1), 2_000)
        selected = lines[start : start + count]
        rendered = "\n".join(f"{line_number}|{line}" for line_number, line in enumerate(selected, start=start + 1))
        self.milestones.add("prepared")
        return self._with_grade(
            {"content": rendered[:MAX_OUTPUT_CHARS], "total_lines": len(lines), "truncated": start + count < len(lines)}
        )

    def write_file(self, path: str, content: str) -> dict[str, Any]:
        target = _safe_path(self.root, path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        self.milestones.add("prepared")
        return self._with_grade({"path": target.relative_to(self.root).as_posix(), "bytes": len(content.encode())})

    def skill_view(self, name: str) -> dict[str, Any]:
        if name not in _SKILLS:
            raise OfficeToolError(f"unknown Office skill: {name}")
        self.milestones.add("prepared")
        return self._with_grade({"name": name, "content": _SKILLS[name]})

    def _with_grade(self, result: dict[str, Any]) -> dict[str, Any]:
        grade = grade_workspace(self.record, self.root)
        output = self.root / str(self.record["output_file"])
        if output.is_file() and output.stat().st_size > 0:
            self.milestones.add("output_created")
        progress_score = max(
            (0.05 if "prepared" in self.milestones else 0.0),
            (0.10 if "executed" in self.milestones else 0.0),
            (0.20 if "output_created" in self.milestones else 0.0),
        )
        return {
            **result,
            "progress_score": progress_score,
            "artifact_score": grade["score"],
            "artifact_issues": grade["issues"],
        }


def run_tool(workspace: OfficeWorkspace, name: str, arguments: dict[str, Any]) -> dict[str, Any]:
    """Dispatch one Hermes-compatible Office tool call."""

    try:
        if name == "execute_code":
            return workspace.execute_code(str(arguments.get("code", "")))
        if name == "read_file":
            return workspace.read_file(
                str(arguments.get("path", "")),
                offset=int(arguments.get("offset", 1)),
                limit=int(arguments.get("limit", 2_000)),
            )
        if name == "write_file":
            return workspace.write_file(str(arguments.get("path", "")), str(arguments.get("content", "")))
        if name == "skill_view":
            return workspace.skill_view(str(arguments.get("name", "")))
        return {"error": f"unknown tool: {name}", "artifact_score": 0.0}
    except (OfficeToolError, OSError, ValueError) as exc:
        return {"error": str(exc), "artifact_score": 0.0}


def _safe_path(root: Path, path: str) -> Path:
    candidate = Path(path)
    if candidate.is_absolute():
        raise OfficeToolError("absolute paths are not allowed")
    resolved = (root / candidate).resolve()
    try:
        resolved.relative_to(root.resolve())
    except ValueError as exc:
        raise OfficeToolError("path escapes the Office workspace") from exc
    return resolved


def _read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=False)
    rows = []
    for sheet in workbook.worksheets:
        rows.append(f"# Sheet: {sheet.title}")
        for values in sheet.iter_rows(values_only=True):
            rows.append("\t".join("" if value is None else str(value) for value in values))
    workbook.close()
    return "\n".join(rows)


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    rows = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table_index, table in enumerate(document.tables, start=1):
        rows.append(f"# Table {table_index}")
        rows.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(rows)


def _read_pdf(path: Path) -> str:
    content = path.read_bytes()
    return "\n".join(
        (
            f"PDF header: {content[:8]!r}",
            f"File size: {len(content)} bytes",
            f"Page markers: {content.count(b'/Type /Page')}",
            f"Has EOF trailer: {content.rstrip().endswith(b'%%EOF')}",
        )
    )


def decode_tool_arguments(raw: str | dict[str, Any] | None) -> dict[str, Any]:
    if isinstance(raw, dict):
        return raw
    try:
        value = json.loads(raw or "{}")
    except json.JSONDecodeError:
        return {}
    return value if isinstance(value, dict) else {}


_HERMES_TOOLS_COMPAT = r'''"""Small compatibility layer used by Office execute_code."""
from pathlib import Path


def _safe(path):
    root = Path.cwd().resolve()
    target = (root / path).resolve()
    target.relative_to(root)
    return target


def write_file(path, content):
    target = _safe(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return {"path": str(target.relative_to(Path.cwd())), "bytes": len(content.encode("utf-8"))}


def read_file(path, offset=1, limit=2000):
    target = _safe(path)
    if target.suffix.lower() == ".xlsx":
        from openpyxl import load_workbook
        workbook = load_workbook(target, read_only=True, data_only=False)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"# Sheet: {sheet.title}")
            lines.extend("\t".join("" if value is None else str(value) for value in row) for row in sheet.iter_rows(values_only=True))
        workbook.close()
    elif target.suffix.lower() == ".docx":
        from docx import Document
        document = Document(target)
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            lines.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    elif target.suffix.lower() == ".pdf":
        content = target.read_bytes()
        lines = [
            f"PDF header: {content[:8]!r}",
            f"File size: {len(content)} bytes",
            f"Page markers: {content.count(b'/Type /Page')}",
            f"Has EOF trailer: {content.rstrip().endswith(b'%%EOF')}",
        ]
    else:
        lines = target.read_text(encoding="utf-8", errors="replace").splitlines()
    selected = lines[max(offset - 1, 0):max(offset - 1, 0) + limit]
    return {"content": "\n".join(selected), "total_lines": len(lines)}


def search_files(pattern, target="content", path=".", file_glob=None, limit=50):
    import re
    base = _safe(path)
    candidates = [base] if base.is_file() else sorted(base.rglob(file_glob or "*"))
    matches = []
    for candidate in candidates:
        if not candidate.is_file() or ".areno_runtime" in candidate.parts:
            continue
        if target == "files":
            if re.search(pattern, candidate.name):
                matches.append({"path": str(candidate.relative_to(Path.cwd()))})
        elif candidate.suffix.lower() not in {".xlsx", ".docx"}:
            for number, line in enumerate(candidate.read_text(encoding="utf-8", errors="replace").splitlines(), 1):
                if re.search(pattern, line):
                    matches.append({"path": str(candidate.relative_to(Path.cwd())), "line": number, "text": line})
        if len(matches) >= limit:
            break
    return {"matches": matches}
'''
